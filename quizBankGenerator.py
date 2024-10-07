import os
import re
import sqlite3
import random
import tempfile

# from fpdf import FPDF
from docx import Document
from reportlab.pdfgen import canvas
from PIL import Image
import matplotlib.pyplot as plt
import os

custom_page_size = (612, 936)


def get_questions_by_subject(subject, limit=10):
    conn = sqlite3.connect('question_bank.db')
    cursor = conn.cursor()

    cursor.execute("SELECT question, option_a, option_b, option_c, option_d, correct_answer FROM questions "
                   "WHERE subject = ? ORDER BY RANDOM() LIMIT ?", (subject, limit))

    questions = cursor.fetchall()
    conn.close()

    return questions


'''
    formatted_questions = []
    for question in questions:
        formatted_questions.append({
            'question': question[0],
            'options': [f'A) {question[1]}', f'B) {question[2]}', f'C) {question[3]}', f'D) {question[4]}'],
            'answer': question[5]
        })

    return formatted_questions


def generate_pdf(test_bank, filename='test_bank.pdf'):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font('Arial', 'B', 10)

    for idx, question in enumerate(test_bank, 1):
        pdf.cell(200, 10, f'Q{idx}, {question["question"]}', ln=True)
        pdf.ln(5)
        for option in question['options']:
            pdf.cell(200, 10, option, ln=True)
        pdf.ln(10)

    pdf.output(filename)
    print(f'PDF generated: {filename}')
'''


def draw_mathtext(c, x, y, text, font_size=12):
    fig = plt.figure(figsize=(0.001, 0.001), dpi=100)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.text(0.5, 0.5, text, fontsize=font_size, horizontalalignment='center', verticalalignment='center')
    ax.axis('off')

    with tempfile.NamedTemporaryFile(suffix='png', delete=False) as tmpfile:
        fig.savefig(tmpfile.name, format='png', bbox_inches='tight', transparent=True)
        plt.close(fig)

        try:
            with Image.open(tmpfile.name) as img:
                width, height = img.size

            c.drawImage(tmpfile.name, x, y, width=width * 0.75, height=height * 0.75)

        except Exception as e:
            print(f"Error opening image: {e}")

    os.remove(tmpfile.name)


def draw_text(c, x, y, text, fontsize=12):
    c.setFont("Helvetica", fontsize)

    max_width = custom_page_size[0] - 100
    lines = []
    words = text.split(' ')
    current_line = ""

    for word in words:
        test_line = f'{current_line} {word}'.strip()
        text_width = c.stringWidth(test_line, "Helvetica", fontsize)

        if text_width <= max_width:
            current_line = test_line
        else:
            lines.append(current_line)
            current_line = word

    if current_line:
        lines.append(current_line)

    for line in lines:
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", fontsize)
            y = custom_page_size[1] - 50

        if contains_latex(line):
            draw_mathtext(c, x, y, line, fontsize)
            y -= 15 * 1.5
        else:
            c.drawString(x, y, line)
            y -= 15

    return y


def contains_latex(text):
    return bool(re.search(r'\$.*?\$|\\\[.*?\\\]', text))


def generate_test_bank_pdf(test_bank, filename='test_bank.pdf'):
    c = canvas.Canvas(filename, pagesize=custom_page_size)
    width, height = custom_page_size

    left_margin = 50
    right_margin = 50
    top_margin = 50
    bottom_margin = 50

    x = left_margin
    y = height - top_margin

    c.setFont("Helvetica", 12)

    for idx, question in enumerate(test_bank, start=1):
        question_text, option_a, option_b, option_c, option_d, correct_answer = question

        y = draw_text(c, x, y, f'Q{idx}. {question_text}', fontsize=12)

        y = draw_text(c, x, y, f'a) {option_a}', fontsize=12)
        y = draw_text(c, x, y, f'b) {option_b}', fontsize=12)
        y = draw_text(c, x, y, f'c) {option_c}', fontsize=12)
        y = draw_text(c, x, y, f'd) {option_d}', fontsize=12)
        y -= 15

        if y < bottom_margin:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - top_margin

    c.save()


def generate_word(test_bank, filename='test_bank.docx'):
    doc = Document()
    doc.add_heading('Test Bank', 0)

    for idx, question in enumerate(test_bank, 1):
        doc.add_paragraph(f'Q{idx}. {question["question"]}')
        for option in question['options']:
            doc.add_paragraph(option, style='List Bullet')
        doc.add_paragraph()
    doc.save(filename)
    print(f'Word document generated: {filename}')


def generate_test_from_db(subjects, num_questions_per_subject=10, output_format='pdf'):
    test_bank = []

    for subject in subjects:
        questions = get_questions_by_subject(subject, num_questions_per_subject)
        test_bank.extend(questions)

    random.shuffle(test_bank)
    print(test_bank)

    if output_format == 'pdf':
        generate_test_bank_pdf(test_bank)
    elif output_format == 'word':
        generate_word(test_bank)
    else:
        print('Unsupported format. Use pdf or word.')


subjects = ['Math', 'Engineering Sciences', 'Professional Electrical Engineering']
generate_test_from_db(subjects, num_questions_per_subject=10, output_format='pdf')
